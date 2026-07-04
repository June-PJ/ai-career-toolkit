#!/usr/bin/env python3
"""Build an ATS-friendly DOCX resume from a Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - user environment guard
    raise SystemExit(
        "python-docx is required. Install it or run with a bundled runtime that includes python-docx."
    ) from exc


ACCENT = RGBColor(35, 91, 150)
ACCENT_DARK = RGBColor(23, 55, 94)
MUTED = RGBColor(92, 99, 112)
TEXT = RGBColor(26, 32, 44)
LIGHT_RULE = "D8DEE8"
SECTION_FILL = "EEF4FB"
BODY_FONT = "Microsoft YaHei"
SERIF_FONT = "SimSun"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=9.5):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_bottom_border(paragraph, color=LIGHT_RULE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_shading(paragraph, fill=SECTION_FILL):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size in [("Heading 1", 22), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT_DARK if style_name != "Heading 1" else RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 2" else 5)
        style.paragraph_format.space_after = Pt(5)


def clean_line(line: str) -> str:
    return line.lstrip("\ufeff").strip().replace("\u00a0", " ")


def split_bold_segments(text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True
        else:
            yield part, False


def add_inline_markdown(paragraph, text: str, size=10.5):
    for chunk, bold in split_bold_segments(text):
        run = paragraph.add_run(chunk)
        set_run_font(run, size=size, bold=bold)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(4)
    add_inline_markdown(p, f" {text}", size=10)


def add_heading(doc: Document, level: int, text: str):
    p = doc.add_paragraph()
    if level == 2:
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        add_shading(p)
        run = p.add_run(f"  {text}")
        set_run_font(run, size=12.3, bold=True, color=ACCENT_DARK)
    else:
        add_project_heading(p, text)


def add_contact(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = 1
    p.paragraph_format.space_after = Pt(2.5)
    add_inline_markdown(p, text, size=9.8)


def add_project_heading(paragraph, text: str):
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2.5)
    if "|" in text:
        title, meta = [part.strip() for part in text.split("|", 1)]
        title_run = paragraph.add_run(title)
        set_run_font(title_run, size=11.2, bold=True, color=TEXT)
        sep_run = paragraph.add_run("  |  ")
        set_run_font(sep_run, size=9.5, color=MUTED)
        meta_run = paragraph.add_run(meta)
        set_run_font(meta_run, size=9.5, color=MUTED)
        return
    run = paragraph.add_run(text)
    set_run_font(run, size=11.2, bold=True, color=TEXT)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_resume_header(doc: Document, name: str, contacts: list[str], photo_path: Path | None = None):
    role, contact_items = normalize_contacts(contacts)

    if photo_path:
        table = doc.add_table(rows=1, cols=2)
        remove_table_borders(table)
        table.autofit = False
        table.columns[0].width = Cm(13.2)
        table.columns[1].width = Cm(3.7)
        left = table.cell(0, 0)
        right = table.cell(0, 1)

        p = left.paragraphs[0]
        run = p.add_run(name)
        set_run_font(run, size=24, bold=True, color=RGBColor(15, 23, 42))
        p.paragraph_format.space_after = Pt(1)

        if role:
            role_p = left.add_paragraph()
            role_p.paragraph_format.space_after = Pt(4)
            role_run = role_p.add_run(role)
            set_run_font(role_run, size=11, bold=True, color=ACCENT)

        if contact_items:
            add_contact_line(left, contact_items[:2], space_after=2)
            add_contact_line(left, contact_items[2:], space_after=6)

        rp = right.paragraphs[0]
        rp.alignment = 2
        photo_run = rp.add_run()
        photo_run.add_picture(str(photo_path), width=Cm(2.55))
        table.rows[0].height = Cm(3.0)
        return

    p = doc.add_paragraph()
    p.alignment = 0
    run = p.add_run(name)
    set_run_font(run, size=24, bold=True, color=RGBColor(15, 23, 42))
    p.paragraph_format.space_after = Pt(1)

    if role:
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_after = Pt(4)
        role_run = role_p.add_run(role)
        set_run_font(role_run, size=11, bold=True, color=ACCENT)

    if contact_items:
        contact = doc.add_paragraph()
        contact.paragraph_format.space_after = Pt(6)
        for index, item in enumerate(contact_items):
            if index:
                sep = contact.add_run("   |   ")
                set_run_font(sep, size=9, color=RGBColor(170, 176, 186))
            run = contact.add_run(item)
            set_run_font(run, size=9.2, color=MUTED)


def add_contact_line(container, items: list[str], space_after=2):
    if not items:
        return
    contact = container.add_paragraph()
    contact.paragraph_format.space_after = Pt(space_after)
    for index, item in enumerate(items):
        if index:
            sep = contact.add_run("   |   ")
            set_run_font(sep, size=9, color=RGBColor(170, 176, 186))
        item_run = contact.add_run(item)
        set_run_font(item_run, size=9.2, color=MUTED)


def extract_header(markdown: str) -> tuple[str, list[str]]:
    name = "简历"
    contacts: list[str] = []
    in_basic_info = False
    for raw in markdown.splitlines():
        line = clean_line(raw)
        if line.startswith("# "):
            name = line[2:].strip() or name
            continue
        if line.startswith("## "):
            in_basic_info = line[3:].strip() == "基本信息"
            continue
        if in_basic_info and line and not line.startswith("#"):
            contacts.append(line)
    return name, contacts


def strip_protocol(text: str) -> str:
    return text.replace("https://", "").replace("http://", "")


def normalize_contacts(contacts: list[str]) -> tuple[str, list[str]]:
    role = ""
    items: list[str] = []
    for contact in contacts:
        for part in [strip_protocol(x.strip()) for x in contact.split("|")]:
            if not part:
                continue
            if "开发" in part or "工程师" in part or "方向" in part:
                role = part
            else:
                items.append(part)
    return role, items


def should_skip(line: str) -> bool:
    return (
        not line
        or line in {"---", "***", "___"}
        or line.startswith("```")
    )


def derive_basename(markdown: str, fallback: str) -> str:
    name = fallback
    target = "简历"
    for line in markdown.splitlines():
        line = clean_line(line)
        if line.startswith("# "):
            name = line[2:].strip() or name
        if "|" in line and ("开发" in line or "工程师" in line or "方向" in line):
            fields = [part.strip() for part in line.split("|")]
            target = fields[-1] or target
            break
    raw = f"{name}-{target}"
    return re.sub(r'[\\/:*?"<>|\s]+', "-", raw).strip("-")


def build_docx(markdown_path: Path, output_path: Path, photo_path: Path | None = None):
    markdown = markdown_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    name, contacts = extract_header(markdown)
    add_resume_header(doc, name, contacts, photo_path)

    in_basic_info = False

    for raw in markdown.splitlines():
        line = clean_line(raw)
        if should_skip(line):
            continue

        if line.startswith("# "):
            in_basic_info = False
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            in_basic_info = title == "基本信息"
            if in_basic_info:
                continue
            add_heading(doc, 2, title)
            continue

        if line.startswith("### "):
            add_heading(doc, 3, line[4:].strip())
            in_basic_info = False
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue

        if in_basic_info:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_markdown(p, line, size=10.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path, help="Path to resume.md")
    parser.add_argument("--outdir", type=Path, default=Path("dist"), help="Output directory")
    parser.add_argument("--basename", help="Output basename without extension")
    parser.add_argument("--photo", type=Path, help="Optional headshot path for resumes that should include a photo")
    return parser.parse_args()


def main():
    args = parse_args()
    markdown_path = args.markdown.resolve()
    if not markdown_path.exists():
        raise SystemExit(f"Markdown file not found: {markdown_path}")
    photo_path = args.photo.resolve() if args.photo else None
    if photo_path and not photo_path.exists():
        raise SystemExit(f"Photo file not found: {photo_path}")

    markdown = markdown_path.read_text(encoding="utf-8")
    basename = args.basename or derive_basename(markdown, markdown_path.stem)
    outdir = args.outdir.resolve()
    docx_path = outdir / f"{basename}.docx"

    build_docx(markdown_path, docx_path, photo_path)
    print(f"DOCX: {docx_path}")


if __name__ == "__main__":
    main()
